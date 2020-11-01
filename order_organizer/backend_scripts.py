import subprocess
import fitz
import re
import os
import docx
import time
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from datetime import datetime
from PIL import Image, ImageDraw, ImageFont
from django.conf import settings
from django.core.files.storage import default_storage
import boto3


def timer(func):
    def wrapper(*args, **kwargs):
        time_before = time.time()
        func(*args, **kwargs)
        time_after = time.time()
        time_elapsed = time_after - time_before
        print(f'TIME ELAPSED DURING FUNCTION EXECUTION: {time_elapsed} seconds')
    return wrapper


def generate_dict_order_info(PDF_filepath):

    full_text = get_full_text_PDF(PDF_filepath)

    client = list(set(extract_data('Client', full_text)))[0]
    order_num = list(set(extract_data('Order', full_text)))[0]
    order_type = extract_data('Order Type', full_text)
    due_date = list(set(extract_data('Due Date', full_text)))[0]
    due_date_obj = datetime(int(due_date.split('-')[2]), int(due_date.split('-')[1]), int(due_date.split('-')[0]), 0, 0, 0, 0)
    date_inserted = datetime.now()
    part_refs = wrapper_func(extract_data('Part Ref', full_text), 'ref')
    part_names = wrapper_func(extract_data('Part Name', full_text), 'name')
    part_qttys = wrapper_func(extract_data('Part Qtty', full_text), 'qtty')
    part_units = wrapper_func(extract_data('Part Unit', full_text), 'unit')
    part_storage = ['4' for _ in range(len(part_refs))]
    part_ready_qtty = ['0' for _ in range(len(part_refs))]
    part_stock_qtty = ['0' for _ in range(len(part_refs))]

    part_list = list(zip(part_storage, part_refs, part_names, part_qttys, part_ready_qtty, part_stock_qtty, part_units))
    parts_string = generate_parts_string(part_list)

    # EPROD 82 = 81 artigos
    dict_order_info = {
        'client': client[0:30], #max lengh = 30 chars
        'order_type': order_type,
        'order_num': order_num,
        'due_date': due_date_obj,
        'date_inserted': date_inserted,
        'pdf_file_path': PDF_filepath,
        'parts_string': parts_string
    }

    return dict_order_info


def get_full_text_PDF(PDF_filepath):
    doc = fitz.Document(PDF_filepath)
    num_pages = doc.pageCount

    full_text = ''

    for num in range(num_pages):
        page = doc.loadPage(num)

        page_text = page.getText('text')

        if 'Original' in page_text:
            full_text += page_text

    return full_text


def extract_data(type, text_to_search):
    if type == 'Order Type':
        if 'Encomenda Interna Produção' in text_to_search:
            return 'EPROD'
        else:
            return 'ENCOM'
    else:
        search_pattern = get_search_pattern(type)

        search = re.compile(search_pattern[0])
        results = search.finditer(text_to_search)

        results_list = []

        for result in results:
            results_list.append(result.group(search_pattern[1]))

        return results_list


def get_search_pattern(type):
    if type == 'Client':
        return r'Exmo\.\(s\) Sr.\(s\)\n(.+)\n', 1
    elif type == 'Order':
        return r'\b(.+)\nNº Requisição', 1
    elif type == 'Due Date':
        return r'Data Req\.\n\d{1,2}-\d{1,2}-\d{4}\n(\d{1,2}-\d{1,2}-\d{4})', 1
        # return r'\d{0,2}-\d{2}-\d{4}\n(\d{0,2})\n(([0-9A-Z]){12}|([0-9A-Z]){8}|1)\n[A-Za-z0-9].+(\n.+)?\n \d+(\.|,)(\d{0,3}(\.|,))?\d{3}\n', 1
    elif type == 'Part Ref':
        # r'(([0-9A-Z].*){8}|([0-9].*){12}|1)\n[A-Za-z0-9].+(\n.+)?\n \d+(\.|,)(\d{0,3}(\.|,))?\d{3}\n', 1
        return r'(([0-9A-Z]){12}|([0-9A-Z]){8}|1)\n[A-Za-z0-9].+(\n.+)?\n \d+(\.|,)(\d{0,3}(\.|,))?\d{3}\n', 1 #Part Ref #r'(\b[A-Z0-9]{12}\n|\b1\n)'
    elif type == 'Part Name':
        # r'\b([A-Za-z].+(\n.+)?)\n \d+(\.|,)(\d{0,3}(\.|,))?\d{3}\n', 1
        return r'(([0-9A-Z]){12}|([0-9A-Z]){8}|1)\n([A-Za-z0-9].+(\n.+)?)\n \d+(\.|,)(\d{0,3}(\.|,))?\d{3}\n', 4 #Part Name
    elif type == 'Part Qtty':
        return r' \d+(\.|,)(\d{0,3}(\.|,))?\d{3}\n', 0 #Part qtty
    elif type == 'Part Unit':
        return r'( \d+(\.|,)(\d{0,3}(\.|,))?\d{3}\n)(UN|KG|MT|un|kg|mt|Un|Kg|Mt|M2|m2)\n', 5 #Part unit


def wrapper_func(list, type_list, target_len=0):
    corrected_list = [elem.replace('\n', ' ') for elem in list]

    if type_list == 'storage':
        pass

    elif type_list == 'qtty':
        corrected_list = [elem.replace(' ', '') for elem in corrected_list]

    return corrected_list


def generate_parts_string(part_list):
    part_string = ''

    part_list = sorted(part_list, key=lambda x: x[2]) #SORTS BY PART NAME

    for part in part_list:
        sub_part_string = '|?!|'.join(part) + '|?!|'

        part_string += sub_part_string + '/?_!/'

    return part_string


def delete_images_folder(path):
    for file in os.listdir(path):
        if file != 'Etiqueta.jpg':
            os.remove(os.path.join(path, file))


def image_generator(images):
    for image in images:
        yield image


def create_order_image(template_image_path, part_info, counter):
    multiline_part_name = part_info['name'][:38] + '\n' + part_info['name'][38:]

    image = Image.open(template_image_path)
    text = ImageDraw.ImageDraw(image, mode='RGB')

    client_text_position = (135, 132)
    order_num_text_position = (447, 322)
    part_description_text_position = (110, 228)
    part_qtty_text_position = (90, 322)
    date_position = (410, 417)

    fontsize = 25
    fontsize_description = 24
    font = ImageFont.truetype('arial.ttf', fontsize)
    font_description = ImageFont.truetype('arial.ttf', fontsize_description)

    text.text(client_text_position, part_info['client'], fill=(0, 0, 0), font=font)
    text.text(order_num_text_position, part_info['order_num'], fill=(0, 0, 0), font=font)
    text.text(part_description_text_position, multiline_part_name, fill=(0, 0, 0), font=font_description)
    text.text(part_qtty_text_position, part_info['quantity'], fill=(0, 0, 0), font=font)
    text.text(date_position, part_info['date'], fill=(0, 0, 0), font=font)

    subprocess.Popen('mkdir imgs', shell=True)
    subprocess.Popen('echo BEFORE SAVING FILE', shell=True)
    subprocess.Popen('ls', shell=True)

    image.save(f'_ODIMG{counter}.jpg')

    subprocess.Popen('echo CHECKING IF FILE WAS SAVED', shell=True)
    subprocess.Popen('ls', shell=True)

    return f'_ODIMG{counter}.jpg'


def generate_order_tags_document(orders_list):
    images = []
    tag_info = []

    current_date = datetime.now()
    current_date_string = f'{current_date.day}/{current_date.month}/{current_date.year}'

    for order in orders_list:
        order_num = order.doc_type + ' ' + order.order_num_doc
        for order_part in order.parts_list:
            part_info = {
                'client': order.client,
                'order_num': order_num,
                'name': order_part[1] + ' - ' + order_part[2],
                'quantity': str(order_part[3]) + ' ' + order_part[6],
                'date': current_date_string
            }

            tag_info.append(part_info)

    counter = 1
    for part_info in tag_info:
        images.append(create_order_image(os.path.join(settings.STATIC_ROOT, 'Etiqueta.jpg').replace('\\', '/'), part_info, counter))
        counter += 1

    MAX_IMAGES_PAGE = 15

    word_doc = Document()
    sections = word_doc.sections
    margin = 0.45

    for section in sections:
        section.top_margin = Cm(margin)
        section.bottom_margin = Cm(margin)
        section.left_margin = Cm(margin)
        section.right_margin = Cm(margin)

        new_height, new_width = section.page_height, section.page_width
        section.orientation = docx.enum.section.WD_ORIENT.PORTRAIT
        section.page_width = new_width
        section.page_height = new_height

    num_empty_order_images = MAX_IMAGES_PAGE - (len(images) % MAX_IMAGES_PAGE)

    p = word_doc.add_paragraph()
    r = p.add_run()

    for _ in range(num_empty_order_images):
        images.append('Etiqueta.jpg')

    image_gen = image_generator(images)
    num_pages = int(len(images) / MAX_IMAGES_PAGE)
    for _ in range(num_pages):
        for __ in range(MAX_IMAGES_PAGE):
            image = next(image_gen)
            r.add_picture(image, width=Cm(6.8), height=Cm(5.08))

        p = word_doc.add_paragraph()
        r = p.add_run()

    try:
        default_storage.delete('Etiquetas_Produção.docx')
    except:
        pass

    word_doc.save('Etiquetas_Produção.docx')

    client = boto3.client('s3')
    client.upload_file('Etiquetas_Produção.docx', 'web-app-logistica-prilux', 'Etiquetas_Produção.docx')


def generate_order_lists_document(orders_list):
    def set_column_widths(table):
        for row in table.rows:
            row.cells[0].width = Cm(0.5)
            row.cells[1].width = Cm(1)
            row.cells[2].width = Cm(9.3)
            row.cells[3].width = Cm(1.5)
            row.cells[4].width = Cm(1.5)
            row.cells[5].width = Cm(1.5)
            row.cells[6].width = Cm(0.7)
            row.cells[7].width = Cm(1.5)

    def vertical_align_text_cells(table):
        for row in table.rows:
            row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row.cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row.cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row.cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row.cells[5].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row.cells[6].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row.cells[7].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def set_table_row_background_color(row):
        pending_quantity = float(row[3].text)
        ready_quantity = float(row[4].text)

        if ready_quantity >= pending_quantity:
            for cell in row:
                shading_elm_2 = parse_xml(r'<w:shd {} w:fill="32CD32"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm_2)

    def return_total_weight(table):
        return '0'

    doc = Document()

    # Setting font size
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(8)

    sections = doc.sections

    for section in sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(0.8)
        section.right_margin = Cm(0.8)

    for order in orders_list:
        order_formatted_due_date = f'{order.due_date.day}/{order.due_date.month}/{order.due_date.year}'

        heading = f"{order.client} | {order.doc_type} nº {order.order_num_doc} | Data Entrega: {order_formatted_due_date}"

        h = doc.add_heading(heading)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        parts_table = doc.add_table(rows=1, cols=8)

        # Setting table style and alignment
        parts_table.style = 'LightGrid'
        parts_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr_cells = parts_table.rows[0].cells

        p = hdr_cells[0].add_paragraph('Arm.')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = hdr_cells[1].add_paragraph('Ref.')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = hdr_cells[2].add_paragraph('Nome')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = hdr_cells[3].add_paragraph('Qtd. Pend')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = hdr_cells[4].add_paragraph('Qtd. Pronta')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = hdr_cells[5].add_paragraph('Qtd. Stock')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = hdr_cells[6].add_paragraph('UN.')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = hdr_cells[7].add_paragraph('Peso Agreg.')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for part in order.parts_list:
            new_row_cells = parts_table.add_row().cells

            p = new_row_cells[0].add_paragraph(part[0])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = new_row_cells[1].add_paragraph(part[1])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = new_row_cells[2].add_paragraph(part[2])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = new_row_cells[3].add_paragraph(str(part[3]))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = new_row_cells[4].add_paragraph(str(part[4]))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = new_row_cells[5].add_paragraph(str(part[5]))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = new_row_cells[6].add_paragraph(part[6])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = new_row_cells[7].add_paragraph('0 kg')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            set_table_row_background_color(new_row_cells)

        set_column_widths(parts_table)
        vertical_align_text_cells(parts_table)

        total_weight = return_total_weight(parts_table)

        p = doc.add_paragraph()
        p = doc.add_paragraph(f'Peso Total: {total_weight} kg')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)

    try:
        default_storage.delete('Listagem.docx')
    except:
        pass

    doc.save('Listagem.docx')

    client = boto3.client('s3')
    client.upload_file('Listagem.docx', 'web-app-logistica-prilux', 'Listagem.docx')


if __name__ == '__main__':
    filepath = 'D:/PROJETOS/PROJETOS PYTHON/Projetos Trabalho/WEB APP LOGÍSTICA/web_app_log/media/pdfs_encomendas/ENCOM 58_9 ARTIGOS.docx'

    dict = generate_dict_order_info(filepath)

    print(dict['parts_string'])

