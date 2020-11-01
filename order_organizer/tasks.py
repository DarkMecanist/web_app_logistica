import subprocess
import fitz
import re
import os
import docx
import time
import boto3
from background_task import background
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from datetime import datetime
from PIL import Image, ImageDraw, ImageFont
from django.conf import settings
from django.core.files.storage import default_storage
from .models import Part


def timer(func):
    def wrapper(*args, **kwargs):
        time_before = time.time()
        func(*args, **kwargs)
        time_after = time.time()
        time_elapsed = time_after - time_before
        print(f'TIME ELAPSED DURING FUNCTION EXECUTION: {time_elapsed} seconds')
    return wrapper


def generate_dict_order_info(PDF_filepath):

    full_text = get_full_text_PDF(PDF_filepath)
    print(full_text)

    client = list(set(extract_data('Client', full_text)))[0]
    order_num = list(set(extract_data('Order', full_text)))[0]
    order_type = extract_data('Order Type', full_text)
    due_date = list(set(extract_data('Due Date', full_text)))[0]
    due_date_obj = datetime(int(due_date.split('-')[2]), int(due_date.split('-')[1]), int(due_date.split('-')[0]), 0, 0, 0, 0)
    date_inserted = datetime.now()
    # part_refs = extract_data('Part Ref', full_text)
    # part_names = extract_data('Part Name', full_text)
    # part_qttys = extract_data('Part Qtty', full_text)
    # part_units = extract_data('Part Unit', full_text)
    part_refs = wrapper_func(extract_data('Part Ref', full_text), 'ref')
    part_names = wrapper_func(extract_data('Part Name', full_text), 'name')
    part_qttys = wrapper_func(extract_data('Part Qtty', full_text), 'qtty')
    part_units = wrapper_func(extract_data('Part Unit', full_text), 'unit')
    part_storage = ['4' for _ in range(len(part_refs))]
    part_ready_qtty = ['0' for _ in range(len(part_refs))]
    part_stock_qtty = ['0' for _ in range(len(part_refs))]

    part_list = list(zip(part_storage, part_refs, part_names, part_qttys, part_ready_qtty, part_stock_qtty, part_units))
    parts_string = generate_parts_string(part_list)
    # parts_string = ''

    # EPROD 82 = 81 artigos
    dict_order_info = {
        'client': client[0:30], #max lengh = 30 chars
        'order_type': order_type,
        'order_num': order_num,
        'due_date': due_date_obj,
        'date_inserted': date_inserted,
        'pdf_file_path': PDF_filepath,
        'parts_string': parts_string
    }

    print(f'DICT ORDERS INFO: {dict_order_info}')

    return dict_order_info


def get_full_text_PDF(PDF_filepath):
    doc = fitz.Document(PDF_filepath)
    num_pages = doc.pageCount

    full_text = ''

    for num in range(num_pages):
        page = doc.loadPage(num)

        page_text = page.getText('text')

        if 'Original' in page_text:
            full_text += page_text

    return full_text


def extract_data(type, text_to_search):
    if type == 'Order Type':
        if 'Encomenda Interna Produção' in text_to_search:
            return 'EPROD'
        else:
            return 'ENCOM'
    else:
        search_pattern = get_search_pattern(type)

        search = re.compile(search_pattern[0])
        results = search.finditer(text_to_search)

        results_list = []

        for result in results:
            results_list.append(result.group(search_pattern[1]))

        print(f'Found {len(results_list)} results for type {type}')
        print(results_list)

        return results_list


def get_search_pattern(type):
    if type == 'Client':
        return r'Exmo\.\(s\) Sr.\(s\)\n(.+)\n', 1
    elif type == 'Order':
        return r'\b(.+)\nNº Requisição', 1
    elif type == 'Due Date':
        return r'Data Req\.\n\d{1,2}-\d{1,2}-\d{4}\n(\d{1,2}-\d{1,2}-\d{4})', 1
    elif type == 'Part Ref':
        return r'(\b[0-9]{8}|\b[A-Z0-9]{12}|\b1)\n(.+|.+\n.+|.+\n.+\n.+)\n(UN|KG|MT|un|kg|mt|Un|Kg|Mt|M2|m2)', 1  # Part Ref #r'(\b[A-Z0-9]{12}\n|\b1\n)'
    elif type == 'Part Name':
        return r'(\b[0-9]{8}|\b[A-Z0-9]{12}|\b1)\n(.+|.+\n.+|.+\n.+\n.+)\n(UN|KG|MT|un|kg|mt|Un|Kg|Mt|M2|m2)', 2  # Part Name r'\b(.+|.+\n.+|.+\n.+\n.+)\nUN'
    elif type == 'Part Qtty':
        return r'\b(UN|KG|MT|un|kg|mt|Un|Kg|Mt|M2|m2|MT2|Mt2|mt2)\n (([0-9]+|[0-9]+\s[0-9]+),[0-9]+)', 2  # Part qtty
    elif type == 'Part Unit':
        return r'\b(UN|KG|MT|un|kg|mt|Un|Kg|Mt|M2|m2|MT2|Mt2|mt2)\n ([0-9]+|[0-9]+\s[0-9]+,[0-9]+)', 1  # Part unit

    # if type == 'Client':
    #     return r'Exmo\.\(s\) Sr.\(s\)\n(.+)\n', 1
    # elif type == 'Order':
    #     return r'\b(.+)\nNº Requisição', 1
    # elif type == 'Due Date':
    #     return r'Data Req\.\n\d{1,2}-\d{1,2}-\d{4}\n(\d{1,2}-\d{1,2}-\d{4})', 1
    #     # return r'\d{0,2}-\d{2}-\d{4}\n(\d{0,2})\n(([0-9A-Z]){12}|([0-9A-Z]){8}|1)\n[A-Za-z0-9].+(\n.+)?\n \d+(\.|,)(\d{0,3}(\.|,))?\d{3}\n', 1
    # elif type == 'Part Ref':
    #     # r'(([0-9A-Z].*){8}|([0-9].*){12}|1)\n[A-Za-z0-9].+(\n.+)?\n \d+(\.|,)(\d{0,3}(\.|,))?\d{3}\n', 1
    #     return r'\d{1,2}\n(([0-9A-Z]){12}|([0-9A-Z]){8}|1)\n[A-Za-z0-9].+(\n.+)?\n \d+(\.|,)(\d{0,3}(\.|,))?\d{3}\n', 1 #Part Ref #r'(\b[A-Z0-9]{12}\n|\b1\n)'
    # elif type == 'Part Name':
    #     # r'\b([A-Za-z].+(\n.+)?)\n \d+(\.|,)(\d{0,3}(\.|,))?\d{3}\n', 1
    #     return r'(([0-9A-Z]){12}|([0-9A-Z]){8}|1)\n([A-Za-z0-9].+(\n.+)?)\n \d+(\.|,)(\d{0,3}(\.|,))?\d{3}\n', 4 #Part Name
    # elif type == 'Part Qtty':
    #     return r' \d+(\.|,)(\d{0,3}(\.|,))?\d{3}\n', 0 #Part qtty
    # elif type == 'Part Unit':
    #     return r'( \d+(\.|,)(\d{0,3}(\.|,))?\d{3}\n)(UN|KG|MT|un|kg|mt|Un|Kg|Mt|M2|m2)\n', 5 #Part unit


def wrapper_func(list, type_list, target_len=0):
    corrected_list = [elem.replace('\n', ' ') for elem in list]

    if type_list == 'storage':
        pass

    elif type_list == 'qtty':
        corrected_list = [elem.replace(' ', '') for elem in corrected_list]

    return corrected_list


def generate_parts_string(part_list):
    part_string = ''

    part_list = sorted(part_list, key=lambda x: x[2]) #SORTS BY PART NAME

    for part in part_list:
        sub_part_string = '|?!|'.join(part) + '|?!|'

        part_string += sub_part_string + '/?_!/'

    return part_string


def delete_images_folder(path):
    for file in os.listdir(path):
        if file != 'Etiqueta.jpg':
            os.remove(os.path.join(path, file))


def image_generator(images):
    for image in images:
        yield image


def create_order_image(template_image_path, part_info, counter):
    multiline_part_name = part_info['name'][:32] + '\n' + part_info['name'][32:]

    image = Image.open(template_image_path)
    text = ImageDraw.ImageDraw(image, mode='RGB')

    client_text_position = (135, 132)
    order_num_text_position = (445, 322)
    part_description_text_position = (110, 228)
    part_qtty_text_position = (90, 322)
    date_position = (410, 417)

    fontsize = 25
    fontsize_description = 24
    font = ImageFont.truetype('arial.ttf', fontsize)
    font_description = ImageFont.truetype('arial.ttf', fontsize_description)

    text.text(client_text_position, part_info['client'], fill=(0, 0, 0), font=font)
    text.text(order_num_text_position, part_info['order_num'], fill=(0, 0, 0), font=font)
    text.text(part_description_text_position, multiline_part_name, fill=(0, 0, 0), font=font_description)
    text.text(part_qtty_text_position, part_info['quantity'], fill=(0, 0, 0), font=font)
    text.text(date_position, part_info['date'], fill=(0, 0, 0), font=font)

    subprocess.Popen('mkdir imgs', shell=True)

    image.save(f'_ODIMG{counter}.jpg')

    return f'_ODIMG{counter}.jpg'


def convert_objects_to_lists(objects_list, model):

    if model == 'order':
        def return_formatted_date(datetime_object):
            return f'{datetime_object.day}/{datetime_object.month}/{datetime_object.year}'

        def convert_parts_string_to_parts_list(parts_string):
            parts_list = [elem.split('|?!|') for elem in parts_string.split('/?_!/')]
            del parts_list[-1]

            for part in parts_list:
                del part[-1]
                for part_info in part:
                    if part.index(part_info) == 3 or part.index(part_info) == 4 or part.index(part_info) == 5:
                        part[part.index(part_info)] = float(part_info.replace('.', '').replace(',', '.'))

            return parts_list

        new_list = [[obj.client, obj.doc_type, obj.order_num_doc, obj.associated_doc_type, obj.order_num_associated_doc,
                     return_formatted_date(obj.due_date), convert_parts_string_to_parts_list(obj.parts_string)] for obj in objects_list]

    elif model == 'shipping':
        def convert_parts_string_to_parts_list(parts_string):
            parts_list = [elem.split('|?!|') for elem in parts_string.split('/?_!/')]
            del parts_list[-1]

            for part in parts_list:
                del part[-1]
                for part_info in part:
                    if part.index(part_info) == 3:
                        part[part.index(part_info)] = float(part_info.replace('.', '').replace(',', '.'))

            return parts_list

        new_list = [[obj.client, obj.doc_type, obj.num_doc, obj.associated_doc_type, obj.order_num_associated_doc, obj.vehicle_info,
                     obj.total_weight, convert_parts_string_to_parts_list(obj.parts_string)] for obj in objects_list]

    elif model == 'part':
        new_list = [[obj.storage, obj.reference, obj.name, obj.stock_quantity, obj.stock_unit, obj.material, obj.weight] for obj in objects_list]

    return new_list


def get_part_weight(part_ref, part_name):
    part_weight = 'N/A'

    try:
        if part_ref == '1':
            part_list = Part.objects.filter(reference=part_ref)

            for part in part_list:
                if part.name == part_name:
                    part_weight = part.weight
        else:
            part = Part.objects.get(reference=part_ref)
            part_weight = part.weight
    except:
        # print(f'PART {part_ref}: {part_name} has no corresponding weight in the database')
        pass

    return part_weight


@background(schedule=0)
def generate_order_tags_document(orders_list):
    images = []
    tag_info = []

    current_date_string = f'{datetime.now().day}/{datetime.now().month}/{datetime.now().year}'

    for order in orders_list:
        order_num = order[1] + ' ' + order[2]
        for order_part in order[6]:
            part_info = {
                'client': order[0],
                'order_num': order_num.split(' ')[1],
                'name': order_part[1] + ' - ' + order_part[2],
                'quantity': str(order_part[3]) + ' ' + order_part[6],
                'date': current_date_string
            }

            tag_info.append(part_info)

    counter = 1
    for part_info in tag_info:
        images.append(create_order_image(os.path.join(settings.STATIC_ROOT, 'Etiqueta.jpg').replace('\\', '/'), part_info, counter))
        counter += 1

    MAX_IMAGES_PAGE = 15

    word_doc = Document()
    sections = word_doc.sections
    margin = 0.45

    for section in sections:
        section.top_margin = Cm(margin)
        section.bottom_margin = Cm(margin)
        section.left_margin = Cm(margin)
        section.right_margin = Cm(margin)

        new_height, new_width = section.page_height, section.page_width
        section.orientation = docx.enum.section.WD_ORIENT.PORTRAIT
        section.page_width = new_width
        section.page_height = new_height

    num_empty_order_images = MAX_IMAGES_PAGE - (len(images) % MAX_IMAGES_PAGE)

    p = word_doc.add_paragraph()
    r = p.add_run()

    for _ in range(num_empty_order_images):
        images.append('Etiqueta.jpg')

    image_gen = image_generator(images)
    num_pages = int(len(images) / MAX_IMAGES_PAGE)
    for _ in range(num_pages):
        for __ in range(MAX_IMAGES_PAGE):
            image = next(image_gen)
            r.add_picture(image, width=Cm(6.8), height=Cm(5.08))

        p = word_doc.add_paragraph()
        r = p.add_run()

    try:
        default_storage.delete('Etiquetas_Produção.docx')
    except:
        pass

    word_doc.save('Etiquetas_Produção.docx')

    client = boto3.client('s3')
    client.upload_file('Etiquetas_Produção.docx', 'web-app-logistica-prilux', 'Etiquetas_Produção.docx')


@background(schedule=0)
def generate_order_lists_document(orders_list):
    def set_column_widths(table):
        for row in table.rows:
            row.cells[0].width = Cm(0.5)
            row.cells[1].width = Cm(1)
            row.cells[2].width = Cm(9.3)
            row.cells[3].width = Cm(1.5)
            row.cells[4].width = Cm(1.5)
            row.cells[5].width = Cm(1.5)
            row.cells[6].width = Cm(0.7)
            row.cells[7].width = Cm(1.5)

    def vertical_align_text_cells(table):
        for row in table.rows:
            row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row.cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row.cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row.cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row.cells[5].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row.cells[6].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row.cells[7].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def set_table_row_background_color(row):
        pending_quantity = float(row[3].text)
        ready_quantity = float(row[4].text)

        if ready_quantity >= pending_quantity:
            for cell in row:
                shading_elm_2 = parse_xml(r'<w:shd {} w:fill="32CD32"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm_2)

    doc = Document()

    # Setting font size
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(8)

    sections = doc.sections

    for section in sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(0.8)
        section.right_margin = Cm(0.8)

    for order in orders_list:
        total_weight = 0

        order_formatted_due_date = order[5]

        heading = f"{order[0]} | {order[1]} nº {order[2]} | Data Entrega: {order_formatted_due_date}"

        h = doc.add_heading(heading)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        parts_table = doc.add_table(rows=1, cols=8)

        # Setting table style and alignment
        parts_table.style = 'LightGrid'
        parts_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr_cells = parts_table.rows[0].cells

        p = hdr_cells[0].add_paragraph('Arm.')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = hdr_cells[1].add_paragraph('Ref.')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = hdr_cells[2].add_paragraph('Nome')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = hdr_cells[3].add_paragraph('Qtd. Pend')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = hdr_cells[4].add_paragraph('Qtd. Pronta')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = hdr_cells[5].add_paragraph('Qtd. Stock')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = hdr_cells[6].add_paragraph('UN.')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = hdr_cells[7].add_paragraph('Peso Agreg.')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for part in order[6]:
            new_row_cells = parts_table.add_row().cells

            p = new_row_cells[0].add_paragraph(part[0])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = new_row_cells[1].add_paragraph(part[1])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = new_row_cells[2].add_paragraph(part[2])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = new_row_cells[3].add_paragraph(str(part[3]))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = new_row_cells[4].add_paragraph(str(part[4]))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = new_row_cells[5].add_paragraph(str(part[5]))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = new_row_cells[6].add_paragraph(part[6])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            part_weight = get_part_weight(part[1], part[2])
            if part_weight != 'N/A':
                part_weight = float(part[4]) * float(part_weight)
                total_weight += part_weight

            p = new_row_cells[7].add_paragraph(str(part_weight))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            set_table_row_background_color(new_row_cells)

        set_column_widths(parts_table)
        vertical_align_text_cells(parts_table)

        p = doc.add_paragraph()
        p = doc.add_paragraph(f'Peso Total Qtd. Pronta: {total_weight} kg')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)

    try:
        default_storage.delete('Listagem.docx')
    except:
        pass

    doc.save('Listagem.docx')

    client = boto3.client('s3')
    client.upload_file('Listagem.docx', 'web-app-logistica-prilux', 'Listagem.docx')

    print('___________________________________\nSUCCESSFULLY GENERATED DOC ORDERS LIST\n___________________________________')


@background(schedule=0)
def generate_shippings_list_document(shippings_list):
    def set_column_widths(table):
        for row in table.rows:
            row.cells[0].width = Cm(0.5)
            row.cells[1].width = Cm(1)
            row.cells[2].width = Cm(9.3)
            row.cells[3].width = Cm(1.5)
            row.cells[4].width = Cm(1.5)
            row.cells[5].width = Cm(1.5)

    def vertical_align_text_cells(table):
        for row in table.rows:
            row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row.cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row.cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row.cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row.cells[5].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc = Document()

    # Setting font size
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(8)

    sections = doc.sections

    for section in sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(0.8)
        section.right_margin = Cm(0.8)

    for shipping in shippings_list:
        heading = f"{shipping[0]} | {shipping[1]} nº {shipping[2]} | Carro: {shipping[5]}"

        h = doc.add_heading(heading)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        parts_table = doc.add_table(rows=1, cols=6)

        # Setting table style and alignment
        parts_table.style = 'LightGrid'
        parts_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr_cells = parts_table.rows[0].cells

        p = hdr_cells[0].add_paragraph('Arm.')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = hdr_cells[1].add_paragraph('Ref.')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = hdr_cells[2].add_paragraph('Nome')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = hdr_cells[3].add_paragraph('Qtd. Exp')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = hdr_cells[4].add_paragraph('UN.')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = hdr_cells[5].add_paragraph('Peso Agreg.')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for part in shipping[7]:
            new_row_cells = parts_table.add_row().cells

            p = new_row_cells[0].add_paragraph(part[0])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = new_row_cells[1].add_paragraph(part[1])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = new_row_cells[2].add_paragraph(part[2])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = new_row_cells[3].add_paragraph(str(part[3]))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = new_row_cells[4].add_paragraph(part[4])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            part_weight = get_part_weight(part[1], part[2])
            if part_weight != 'N/A':
                part_weight = float(part[3]) * float(part_weight)

            p = new_row_cells[5].add_paragraph(str(part_weight))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        set_column_widths(parts_table)
        vertical_align_text_cells(parts_table)

        total_weight = shipping[6]

        p = doc.add_paragraph()
        p = doc.add_paragraph(f'Peso Total: {total_weight}')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)

    try:
        default_storage.delete('Listagem_Expedicoes.docx')
    except:
        pass

    doc.save('Listagem_Expedicoes.docx')

    client = boto3.client('s3')
    client.upload_file('Listagem_Expedicoes.docx', 'web-app-logistica-prilux', 'Listagem_Expedicoes.docx')

    print('___________________________________\nSUCCESSFULLY GENERATED DOC SHIPPINGS LIST\n___________________________________')


@background(schedule=0)
def generate_parts_list_document(parts_list):
    def set_column_widths(table):
        for row in table.rows:
            row.cells[0].width = Cm(0.5)
            row.cells[1].width = Cm(1)
            row.cells[2].width = Cm(9.3)
            row.cells[3].width = Cm(1.5)
            row.cells[4].width = Cm(1)
            row.cells[5].width = Cm(1.5)
            row.cells[6].width = Cm(0.7)

    def vertical_align_text_cells(table):
        for row in table.rows:
            row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row.cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row.cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row.cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row.cells[5].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row.cells[6].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc = Document()

    # Setting font size
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(8)

    sections = doc.sections

    for section in sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(0.8)
        section.right_margin = Cm(0.8)

        current_time = datetime.now()

        heading = f"Listagem Stocks Dia {current_time.day}/{current_time.month}/{current_time.year} - {current_time.hour}:{current_time.minute}H"

        h = doc.add_heading(heading)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        parts_table = doc.add_table(rows=1, cols=7)

        # Setting table style and alignment
        parts_table.style = 'LightGrid'
        parts_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr_cells = parts_table.rows[0].cells

        p = hdr_cells[0].add_paragraph('Arm.')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = hdr_cells[1].add_paragraph('Ref.')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = hdr_cells[2].add_paragraph('Nome')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = hdr_cells[3].add_paragraph('Qtd. Stock')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = hdr_cells[4].add_paragraph('UN.')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = hdr_cells[5].add_paragraph('Material.')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = hdr_cells[6].add_paragraph('Peso Unit.')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for part in parts_list:
            new_row_cells = parts_table.add_row().cells

            p = new_row_cells[0].add_paragraph(part[0])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = new_row_cells[1].add_paragraph(part[1])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = new_row_cells[2].add_paragraph(part[2])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = new_row_cells[3].add_paragraph(str(part[3]))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = new_row_cells[4].add_paragraph(part[4])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = new_row_cells[5].add_paragraph(part[5])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = new_row_cells[6].add_paragraph(str(part[6]))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        set_column_widths(parts_table)
        vertical_align_text_cells(parts_table)

    try:
        default_storage.delete('Listagem_Artigos.docx')
    except:
        pass

    doc.save('Listagem_Artigos.docx')

    client = boto3.client('s3')
    client.upload_file('Listagem_Artigos.docx', 'web-app-logistica-prilux', 'Listagem_Artigos.docx')

    print('___________________________________\nSUCCESSFULLY GENERATED DOC PARTS LIST\n___________________________________')

